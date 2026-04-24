"""
👖 BLUE JEANS WEB NOVEL ENGINE v2.0 — parser.py
기획서 파일 파싱 모듈 (DOCX / HWP / PDF / TXT)
© 2026 BLUE JEANS PICTURES
"""

import io
import struct


def parse_docx(file_bytes):
    """DOCX 파일에서 텍스트 추출."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # 테이블
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX 파싱 실패: {e}]"


def parse_hwp(file_bytes):
    """HWP 파일(구버전 5.0)에서 텍스트 추출. olefile + PrvText 스트림 활용."""
    try:
        import olefile
        import zlib

        ole = olefile.OleFileIO(io.BytesIO(file_bytes))

        # 1차: PrvText 스트림 (미리보기 텍스트)
        prvtext = ""
        if ole.exists("PrvText"):
            try:
                raw = ole.openstream("PrvText").read()
                prvtext = raw.decode("utf-16-le", errors="replace")
            except Exception:
                prvtext = ""

        # 2차: BodyText/Section0 스트림 (본문 텍스트)
        body_parts = []
        if ole.exists(["BodyText", "Section0"]):
            try:
                # FileHeader에서 압축 여부 확인
                is_compressed = False
                if ole.exists("FileHeader"):
                    header = ole.openstream("FileHeader").read()
                    flags = struct.unpack_from("<I", header, 36)[0]
                    is_compressed = bool(flags & 1)

                data = ole.openstream(["BodyText", "Section0"]).read()
                if is_compressed:
                    data = zlib.decompress(data, -15)

                i = 0
                while i < len(data):
                    if i + 4 > len(data):
                        break
                    h = struct.unpack_from("<I", data, i)[0]
                    tag_id = h & 0x3FF
                    size = (h >> 20) & 0xFFF
                    if size == 0xFFF:
                        if i + 8 > len(data):
                            break
                        size = struct.unpack_from("<I", data, i + 4)[0]
                        i += 8
                    else:
                        i += 4

                    if tag_id == 67:  # HWPTAG_PARA_TEXT
                        txt = ""
                        j = 0
                        while j < size:
                            if j + 2 > size:
                                break
                            ch = struct.unpack_from("<H", data, i + j)[0]
                            if ch >= 32:
                                txt += chr(ch)
                            elif ch in (13, 10):
                                txt += "\n"
                            j += 2
                        if txt.strip():
                            body_parts.append(txt.strip())
                    i += size
            except Exception:
                pass

        ole.close()

        body_text = "\n".join(body_parts)
        # body_text가 충분하면 사용, 아니면 PrvText fallback
        if len(body_text) >= len(prvtext):
            return body_text
        return prvtext
    except Exception as e:
        return f"[HWP 파싱 실패: {e}]"


def parse_pdf(file_bytes):
    """PDF 파일에서 텍스트 추출. pypdf 또는 pdfplumber."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                parts.append(text)
        return "\n".join(parts)
    except ImportError:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text)
            return "\n".join(parts)
        except Exception as e:
            return f"[PDF 파싱 실패: pypdf/PyPDF2 필요 — {e}]"
    except Exception as e:
        return f"[PDF 파싱 실패: {e}]"


def parse_txt(file_bytes):
    """TXT 파일에서 텍스트 추출 (UTF-8, CP949 자동 시도)."""
    for encoding in ["utf-8", "cp949", "euc-kr", "utf-16"]:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def parse_brief(uploaded_file):
    """
    Streamlit UploadedFile 객체를 받아 텍스트 추출.
    파일 확장자에 따라 적절한 파서 호출.
    """
    if uploaded_file is None:
        return ""

    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()

    if filename.endswith(".docx"):
        return parse_docx(file_bytes)
    elif filename.endswith(".hwp"):
        return parse_hwp(file_bytes)
    elif filename.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename.endswith(".txt") or filename.endswith(".md"):
        return parse_txt(file_bytes)
    else:
        return f"[지원하지 않는 파일 형식: {filename}]"
