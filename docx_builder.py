"""
👖 BLUE JEANS WEB NOVEL ENGINE v2.6 — docx_builder.py
DOCX 출력 모듈
- 회차별 DOCX (수정 편의)
- 시즌 전체 DOCX (피칭/보관용)
- 기획서 DOCX (컨셉 + 캐릭터 + 아크)
© 2026 BLUE JEANS PICTURES
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════
# 디자인 상수
# ══════════════════════════════════════════════
NAVY = RGBColor(0x19, 0x19, 0x70)
YELLOW = RGBColor(0xFF, 0xCB, 0x05)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
DIM = RGBColor(0x8E, 0x8E, 0x99)
LIGHT_BG = "EEEEF6"
YELLOW_BG = "FFCB05"

BODY_FONT = "맑은 고딕"  # 범용. 나눔고딕으로도 변경 가능
HEADING_FONT = "맑은 고딕"


# ══════════════════════════════════════════════
# 공통 유틸
# ══════════════════════════════════════════════
def _set_cell_bg(cell, color_hex):
    """셀 배경색 설정."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_run(paragraph, text, font=BODY_FONT, size=11, bold=False,
             color=None, italic=False):
    """스타일링된 run 추가."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # 한글 폰트 설정
    r = run._element
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    return run


def _section_header_yellow(doc, ko_text, en_text=""):
    """노란색 섹션 헤더 (블루진 스타일)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, YELLOW_BG)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(para, ko_text, font=HEADING_FONT, size=12, bold=True, color=NAVY)
    if en_text:
        _add_run(para, f"   {en_text}", font=HEADING_FONT, size=9, bold=False, color=NAVY)


def _horizontal_rule(doc):
    """수평선."""
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "191970")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _meta_line(doc, label, value):
    """메타 정보 한 줄 (라벨 + 값)."""
    para = doc.add_paragraph()
    _add_run(para, f"{label}: ", size=9, bold=True, color=DIM)
    _add_run(para, str(value), size=9, color=DARK)


# ══════════════════════════════════════════════
# 1. 회차별 DOCX (수정 편의)
# ══════════════════════════════════════════════
def build_episode_docx(episode_text, episode_number, concept=None,
                       plot=None, rating="19", platform="카카오페이지"):
    """
    단일 회차 DOCX 생성 — 수정 편의 중심.
    
    Args:
        episode_text: 본문 (첫 줄에 제목 포함)
        episode_number: 회차 번호
        concept: 컨셉 카드 (optional)
        plot: 회차 플롯 (optional)
        rating: '19' or '15'
        platform: 플랫폼명
    
    Returns:
        bytes — DOCX 파일 바이너리
    """
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 본문 첫 줄에서 제목 추출
    lines = episode_text.strip().split("\n")
    if lines and lines[0].strip():
        title_line = lines[0].strip()
        body_start = 1
        # 빈 줄 건너뛰기
        while body_start < len(lines) and not lines[body_start].strip():
            body_start += 1
        body = "\n".join(lines[body_start:])
    else:
        title_line = f"EP{episode_number}"
        body = episode_text

    # ── 제목 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(title_para, title_line, font=HEADING_FONT, size=20, bold=True, color=NAVY)

    # ── 메타 정보 ──
    char_count = len(body.replace(" ", "").replace("\n", ""))
    total_count = len(body)
    meta_items = []
    if concept:
        meta_items.append(f"장르: {concept.get('genre', '')}")
    meta_items.append(f"분량: {total_count:,}자 (공백 포함) · {char_count:,}자 (공백 제외)")
    meta_items.append(f"수위: {rating}금")
    meta_items.append(f"플랫폼: {platform}")
    if plot:
        cliff = plot.get("cliffhanger", {})
        if cliff.get("type"):
            meta_items.append(f"클리프행어: {cliff.get('type')}")

    meta_para = doc.add_paragraph()
    _add_run(meta_para, " | ".join(meta_items), size=9, color=DIM)

    _horizontal_rule(doc)

    # ── 본문 ──
    for line in body.split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.6
        if line.strip():
            _add_run(para, line, size=11, color=DARK)

    # ── 푸터 (플롯 정보가 있으면) ──
    if plot:
        doc.add_page_break()
        _section_header_yellow(doc, "회차 메타", "EPISODE META")

        opening = plot.get("opening", {})
        if opening:
            _meta_line(doc, "첫 문장", opening.get("hook_line", ""))
            _meta_line(doc, "리캡 방식", opening.get("recap_method", ""))
            _meta_line(doc, "독자 질문", opening.get("question", ""))

        cliff = plot.get("cliffhanger", {})
        if cliff:
            _meta_line(doc, "클리프행어 유형", cliff.get("type", ""))
            _meta_line(doc, "클리프행어 내용", cliff.get("content", ""))
            _meta_line(doc, "다음 회차 연결", cliff.get("next_ep_connection", ""))

    # 저장
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════
# 2. 시즌 전체 DOCX
# ══════════════════════════════════════════════
def build_season_docx(episodes_dict, concept, rating="19", platform="카카오페이지"):
    """
    시즌 전체 DOCX 생성 — 피칭/보관용.
    
    Args:
        episodes_dict: {ep_num: episode_text, ...}
        concept: 컨셉 카드
        rating: '19' or '15'
    
    Returns:
        bytes
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 커버 페이지 ──
    for _ in range(3):
        doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(tagline, "BLUE JEANS PICTURES", font=HEADING_FONT, size=10, bold=True, color=NAVY)

    for _ in range(2):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cover_title, concept.get("title", ""), font=HEADING_FONT, size=32, bold=True, color=NAVY)

    for _ in range(2):
        doc.add_paragraph()

    logline_para = doc.add_paragraph()
    logline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(logline_para, concept.get("logline", ""), size=12, italic=True, color=DARK)

    for _ in range(5):
        doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(meta_para, concept.get("genre", ""), size=11, color=NAVY)
    _add_run(meta_para, "   |   ", size=11, color=DIM)
    _add_run(meta_para, f"{len(episodes_dict)}화 · {rating}금", size=11, color=NAVY)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(date_para, datetime.now().strftime("%Y.%m.%d"), size=10, color=DIM)

    doc.add_page_break()

    # ── 목차 ──
    _section_header_yellow(doc, "목차", "CONTENTS")
    doc.add_paragraph()

    for ep_num in sorted(episodes_dict.keys()):
        text = episodes_dict[ep_num]
        first_line = text.strip().split("\n")[0] if text.strip() else f"EP{ep_num}"
        toc_para = doc.add_paragraph()
        _add_run(toc_para, f"EP{ep_num:03d}", size=10, bold=True, color=NAVY)
        _add_run(toc_para, f"   {first_line}", size=10, color=DARK)

    doc.add_page_break()

    # ── 본문 ──
    for ep_num in sorted(episodes_dict.keys()):
        text = episodes_dict[ep_num]
        lines = text.strip().split("\n")

        if lines and lines[0].strip():
            title_line = lines[0].strip()
            body_start = 1
            while body_start < len(lines) and not lines[body_start].strip():
                body_start += 1
            body = "\n".join(lines[body_start:])
        else:
            title_line = f"EP{ep_num}"
            body = text

        # 회차 제목
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_run(title_para, title_line, font=HEADING_FONT, size=18, bold=True, color=NAVY)

        # 본문
        for line in body.split("\n"):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.6
            if line.strip():
                _add_run(para, line, size=11, color=DARK)

        # 회차 구분
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════
# 3. 기획서 DOCX (컨셉 + 캐릭터 + 아크)
# ══════════════════════════════════════════════
def build_proposal_docx(concept, character_bible=None, core_arc=None,
                        plant_map=None):
    """
    기획서 DOCX 생성 — 투자 피칭용.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 커버 ──
    for _ in range(3):
        doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(company, "BLUE JEANS PICTURES", font=HEADING_FONT, size=10, bold=True, color=NAVY)

    for _ in range(2):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cover_title, concept.get("title", ""), font=HEADING_FONT, size=36, bold=True, color=NAVY)

    doc.add_paragraph()

    sub_title = doc.add_paragraph()
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub_title, "작품 기획서", font=HEADING_FONT, size=14, color=DIM)

    for _ in range(3):
        doc.add_paragraph()

    logline_para = doc.add_paragraph()
    logline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(logline_para, concept.get("logline", ""), size=12, italic=True, color=DARK)

    doc.add_page_break()

    # ── 개요 ──
    _section_header_yellow(doc, "개요", "OVERVIEW")

    _meta_line(doc, "제목", concept.get("title", ""))
    _meta_line(doc, "장르", concept.get("genre", ""))
    _meta_line(doc, "형식", "웹소설")
    tags = concept.get("formula_tags", [])
    if tags:
        _meta_line(doc, "태그", " · ".join(tags))
    config = concept.get("serial_config", {})
    if config:
        _meta_line(doc, "Core Arc", f"{config.get('core_eps', 50)}화")
        _meta_line(doc, "수위", config.get("rating_mode", ""))
        _meta_line(doc, "과금 전환", f"EP{config.get('paywall_ep', 25)}")

    doc.add_paragraph()

    # ── 로그라인 ──
    _section_header_yellow(doc, "로그라인", "LOGLINE")
    para = doc.add_paragraph()
    _add_run(para, concept.get("logline", ""), size=11, color=DARK)

    doc.add_paragraph()

    # ── 시놉시스 ──
    _section_header_yellow(doc, "시놉시스", "SYNOPSIS")
    for para_text in concept.get("synopsis", "").split("\n"):
        if para_text.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            _add_run(para, para_text, size=11, color=DARK)

    doc.add_paragraph()

    # ── 시즌 질문 ──
    _section_header_yellow(doc, "시즌 질문", "SEASON QUESTIONS")
    sq = concept.get("season_questions", {})
    _meta_line(doc, "로맨스 축", sq.get("romance", ""))
    _meta_line(doc, "미스터리 축", sq.get("mystery", ""))

    doc.add_paragraph()

    # ── 세계관 ──
    _section_header_yellow(doc, "세계관", "WORLD")
    para = doc.add_paragraph()
    _add_run(para, concept.get("world", ""), size=11, color=DARK)

    doc.add_paragraph()

    # ── 캐릭터 ──
    _section_header_yellow(doc, "캐릭터", "CHARACTERS")

    # 주인공
    p = concept.get("protagonist", {})
    if p.get("name"):
        para = doc.add_paragraph()
        _add_run(para, "▶ 주인공 ", font=HEADING_FONT, size=12, bold=True, color=NAVY)
        _add_run(para, p.get("name", ""), font=HEADING_FONT, size=12, bold=True, color=DARK)
        _meta_line(doc, "Goal", p.get("goal", ""))
        _meta_line(doc, "Need", p.get("need", ""))
        _meta_line(doc, "Fatal Flaw", p.get("fatal_flaw", ""))
        doc.add_paragraph()

    # 상대역
    for i, li in enumerate(concept.get("love_interests", []), 1):
        para = doc.add_paragraph()
        _add_run(para, f"▶ 상대역 {i} ", font=HEADING_FONT, size=12, bold=True, color=NAVY)
        _add_run(para, f"{li.get('name', '')} ({li.get('role', '')})",
                 font=HEADING_FONT, size=12, bold=True, color=DARK)
        _meta_line(doc, "매력", li.get("appeal", ""))
        _meta_line(doc, "갈등", li.get("conflict", ""))
        doc.add_paragraph()

    # 빌런
    v = concept.get("villain", {})
    if v.get("name"):
        para = doc.add_paragraph()
        _add_run(para, "▶ 빌런 ", font=HEADING_FONT, size=12, bold=True, color=NAVY)
        _add_run(para, f"{v.get('name', '')} ({v.get('role', '')})",
                 font=HEADING_FONT, size=12, bold=True, color=DARK)
        _meta_line(doc, "Wants", v.get("wants", ""))
        _meta_line(doc, "Justification", v.get("justification", ""))
        _meta_line(doc, "Limits", v.get("limits", ""))
        _meta_line(doc, "Win Rate", v.get("win_rate", ""))
        doc.add_paragraph()

    # ── Core Arc 요약 ──
    if core_arc:
        doc.add_page_break()
        _section_header_yellow(doc, "Core Arc", "STRUCTURE")

        for block in core_arc:
            para = doc.add_paragraph()
            _add_run(para, f"[블록 {block.get('block_no', '')}] ",
                     font=HEADING_FONT, size=11, bold=True, color=NAVY)
            _add_run(para, f"{block.get('ep_range', '')} · {block.get('phase', '')}",
                     size=11, bold=True, color=DARK)

            theme_para = doc.add_paragraph()
            _add_run(theme_para, block.get("theme", ""), size=10, color=DIM)

            for ep in block.get("episodes", []):
                ep_para = doc.add_paragraph()
                _add_run(ep_para, f"  EP{ep.get('ep', ''):03d}. ",
                         size=10, bold=True, color=NAVY)
                _add_run(ep_para, ep.get("title", ""),
                         size=10, bold=True, color=DARK)
                _add_run(ep_para, f"  — {ep.get('summary', '')}",
                         size=10, color=DARK)
                _add_run(ep_para, f"  [{ep.get('cliffhanger_type', '')}]",
                         size=9, italic=True, color=DIM)
            doc.add_paragraph()

    # ── 떡밥 맵 ──
    if plant_map and plant_map.get("plants"):
        doc.add_page_break()
        _section_header_yellow(doc, "떡밥 맵", "PLANT & PAYOFF")

        for p in plant_map["plants"]:
            para = doc.add_paragraph()
            _add_run(para, f"🌱 {p.get('name', '')} ",
                     font=HEADING_FONT, size=11, bold=True, color=NAVY)
            _add_run(para, f"({p.get('type', '')})",
                     size=9, italic=True, color=DIM)

            _meta_line(doc, "설명", p.get("description", ""))
            _meta_line(doc, "심기", f"EP{p.get('plant_ep', '')}")
            hints = p.get("hints", [])
            if hints:
                _meta_line(doc, "힌트", ", ".join([f"EP{h}" for h in hints]))
            _meta_line(doc, "회수", f"EP{p.get('payoff_ep', '')}")
            doc.add_paragraph()

    # ── 푸터 ──
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(footer, "© 2026 BLUE JEANS PICTURES", size=9, color=DIM)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
